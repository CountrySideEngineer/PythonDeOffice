import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.section import Sections

def SetFooterLeft(path:str, footers:list) -> None:
	"""Set footer of left.

	Args:
		path(str): Path to file of docx to set footer.
		footers(list):	Collection of strings to set into footer
						One item is one line in footer.
						All items will be joined by "change line" code.
	"""
	align_func = SetFooterAlignmentLeft
	FooterFuncPointer(path=path, footers=footers, align_func=align_func)

def SetFooterCenter(path:str, footers:list) -> None:
	"""Set footer of center.

	Args:
		path(str): Path to file of docx to set footer.
		footers(list):	Collection of strings to set into footer
						One item is one line in footer.
						All items will be joined by "change line" code.
	"""
	align_func = SetFooterAlignmentCenter
	FooterFuncPointer(path=path, footers=footers, align_func=align_func)

def SetFooterRight(path:str, footers:list) -> None:
	"""Set footer of right.

	Args:
		path(str): Path to file of docx to set footer.
		footers(list):	Collection of strings to set into footer
						One item is one line in footer.
						All items will be joined by "change line" code.
	"""
	align_func = SetFooterAlignmentRight
	FooterFuncPointer(path=path, footers=footers, align_func=align_func)

def FooterFuncPointer(path:str, footers:list, align_func) -> None:
	"""Set footer of word file.

	Args:
		path(str):	Path to file to set footer.
		footers(list):	Collection of strings to set into footer
						One item is one line in footer.
						All items will be joined by "change line" code.
		align_func:	Function object to set alignment.
	
	"""
	document = Document(path)
	sections = document.sections
	section_footer = sections[0].footer

	paragraph_sections = document.paragraphs

	RemoveFooter(footer_section=section_footer)

	for footer_item in footers:
		new_pargraph = section_footer.add_paragraph(footer_item)
		align_func(new_pargraph)

	document.save(path)

def SetFooterAlignmentLeft(paragrah) -> None:
	"""Set alignment left to paragraph
	
	Args:
		paragraph:	Paragraph to set alignment.
	"""
	paragrah.alignment = WD_ALIGN_PARAGRAPH.LEFT

def SetFooterAlignmentCenter(paragrah) -> None:
	"""Set alignment center to paragraph
	
	Args:
		paragraph:	Paragraph to set alignment.
	"""
	paragrah.alignment = WD_ALIGN_PARAGRAPH.CENTER

def SetFooterAlignmentRight(paragrah) -> None:
	"""Set alignment right to paragraph
	
	Args:
		paragraph:	Paragraph to set alignment.
	"""
	paragrah.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def RemoveFooter(footer_section:Sections) -> None:
	"""Remove hedaer in document.
	
	Args:
		footer_section(Sections): Sections obeject of footer.
	"""
	paragraphs = footer_section.paragraphs
	for paragraph in paragraphs:
		element = paragraph._element
		element.getparent().remove(element)
		paragraph._p = paragraph._element = None

if '__main__' == __name__:
	path = sys.argv[1]

	footers = [
		'footer_code_011',
		'footer_code_044',
		]

	SetFooterRight(path, footers=footers)
