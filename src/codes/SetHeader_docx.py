import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.section import Sections

def SetHeaderLeft(path:str, headers:list) -> None:
	"""Set header of left.

	Args:
		path(str): Path to file of docx to set header.
		headers(list):	Collection of strings to set into header
						One item is one line in header.
						All items will be joined by "change line" code.
	"""
	align_func = SetHeaderAlignmentLeft
	HeaderFuncPointer(path=path, headers=headers, align_func=align_func)

def SetHeaderCenter(path:str, headers:list) -> None:
	"""Set header of center.

	Args:
		path(str): Path to file of docx to set header.
		headers(list):	Collection of strings to set into header
						One item is one line in header.
						All items will be joined by "change line" code.
	"""
	align_func = SetHeaderAlignmentCenter
	HeaderFuncPointer(path=path, headers=headers, align_func=align_func)

def SetHeaderRight(path:str, headers:list) -> None:
	"""Set header of right.

	Args:
		path(str): Path to file of docx to set header.
		headers(list):	Collection of strings to set into header
						One item is one line in header.
						All items will be joined by "change line" code.
	"""
	align_func = SetHeaderAlignmentRight
	HeaderFuncPointer(path=path, headers=headers, align_func=align_func)

def HeaderFuncPointer(path:str, headers:list, align_func) -> None:
	"""Set header of word file.

	Args:
		path(str):	Path to file to set header.
		headers(list):	Collection of strings to set into header
						One item is one line in header.
						All items will be joined by "change line" code.
		align_func:	Function object to set alignment.
	
	"""
	document = Document(path)
	sections = document.sections
	section_header = sections[0].header

	RemoveHeader(header_section=section_header)

	for header_item in headers:
		new_pargraph = section_header.add_paragraph(header_item)
		align_func(new_pargraph)

	document.save(path)

def SetHeaderAlignmentLeft(paragrah) -> None:
	"""Set alignment left to paragraph
	
	Args:
		paragraph:	Paragraph to set alignment.
	"""
	paragrah.alignment = WD_ALIGN_PARAGRAPH.LEFT

def SetHeaderAlignmentCenter(paragrah) -> None:
	"""Set alignment center to paragraph
	
	Args:
		paragraph:	Paragraph to set alignment.
	"""
	paragrah.alignment = WD_ALIGN_PARAGRAPH.CENTER

def SetHeaderAlignmentRight(paragrah) -> None:
	"""Set alignment right to paragraph
	
	Args:
		paragraph:	Paragraph to set alignment.
	"""
	paragrah.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def RemoveHeader(header_section:Sections) -> None:
	"""Remove hedaer in document.
	
	Args:
		header_section(Sections): Sections obeject of header.
	"""
	paragraphs = header_section.paragraphs
	for paragraph in paragraphs:
		element = paragraph._element
		element.getparent().remove(element)
		paragraph._p = paragraph._element = None

if '__main__' == __name__:
	path = sys.argv[1]

	headers = [
		'header_code_011',
		'header_code_044',
		]

	SetHeaderLeft(path, headers=headers)
